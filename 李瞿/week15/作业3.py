import asyncio
import os
os.environ["OPENAI_API_KEY"] = "sk-facd0ca4f5ae4fada1706bf3938b69d9"
os.environ["OPENAI_BASE_URL"] = "https://dashscope.aliyuncs.com/compatible-mode/v1"

import json
import requests
import urllib.parse
from typing import List, Dict, Any

# 尝试导入 docx，如果失败则提供替代方案
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    print("警告: 无法导入 python-docx，将无法生成 DOCX 文件")
    DOCX_AVAILABLE = False

# 假设以下导入能够正常工作，它们通常来自 agents 库
from agents import Agent, function_tool, AsyncOpenAI, OpenAIChatCompletionsModel, ModelSettings, Runner, \
    set_default_openai_api, set_tracing_disabled

set_default_openai_api("chat_completions")
set_tracing_disabled(True)

MODEL_NAME = "qwen-max"
API_KEY = os.getenv("OPENAI_API_KEY", "sk-facd0ca4f5ae4fada1706bf3938b69d9")
BASE_URL = os.getenv("OPENAI_BASE_URL", "https://dashscope.aliyuncs.com/compatible-mode/v1")

# 初始化 AsyncOpenAI 客户端
llm_client = AsyncOpenAI(
    api_key=API_KEY,
    base_url=BASE_URL
)

# 定义模型设置
model_settings = ModelSettings(
    model=MODEL_NAME,
    client=llm_client,
    temperature=0.3
)

# --- 2. 外部工具（Jina Search & Crawl） ---
# JINA_API_KEY = "jina_8918effb420d4bff8530c9d9f3bbe536NWhiCZdKQFNgoFLd4aganV1XnsaA"
JINA_API_KEY = "jina_96bf4c31c3e440e9917304d3296a4ae9eJs6KD2-zKKnwqvhM4uKmMzXw14R"

def search_jina(query: str) -> str:
    """通过jina进行谷歌搜索，返回JSON格式的搜索结果字符串"""
    print(f"-> [Jina Search] 正在搜索: {query[:50]}...")
    try:
        encoded_query = urllib.parse.quote(query)
        url = f"https://s.jina.ai/?q={encoded_query}&hl=zh-cn"
        headers = {
            "Accept": "application/json", # 返回结果
            "Authorization": f"Bearer {JINA_API_KEY}",
            "X-Respond-With": "no-content"
        }
        response = requests.get(url, headers=headers, timeout=15)
        response.raise_for_status()  # 抛出 HTTP 错误

        # Jina Search 返回的是一个包含结果的 JSON 结构，提取关键信息
        results = response.json().get('data', [])

        # 提取标题、链接和摘要
        formatted_results = []
        for res in results:
            formatted_results.append({
                "title": res.get("title", ""),
                "url": res.get("url", ""),
                "snippet": res.get("content", "")
            })

        return json.dumps(formatted_results, ensure_ascii=False)
    except requests.exceptions.RequestException as e:
        print(f"Error during Jina Search: {e}")
        return json.dumps({"error": str(e), "query": query}, ensure_ascii=False)
    except Exception as e:
        print(f"Unexpected error in Jina Search: {e}")
        return json.dumps({"error": str(e), "query": query}, ensure_ascii=False)


def crawl_jina(url: str) -> str:
    """通过jina抓取完整网页内容，返回Markdown格式的文本"""
    print(f"-> [Jina Crawl] 正在抓取: {url[:50]}...")
    try:
        # Jina Reader API
        headers = {
            "Accept": "application/json",
            "Authorization": f"Bearer {JINA_API_KEY}",
            "X-Respond-With": "content",  # 请求返回完整内容
            "X-Content-Type": "markdown"  # 请求返回 Markdown 格式
        }
        # 使用 r.jina.ai 作为代理
        response = requests.get("https://r.jina.ai/" + url, headers=headers, timeout=20)
        response.raise_for_status()

        # 返回内容通常在 'data' 字段的 'content' 中
        content = response.json().get("data", {}).get("content", f"无法抓取 URL: {url} 的内容。")

        return content
    except requests.exceptions.RequestException as e:
        print(f"Error during Jina Crawl for {url}: {e}")
        return f"抓取失败: {e}"
    except Exception as e:
        print(f"Unexpected error in Jina Crawl for {url}: {e}")
        return f"抓取失败: {e}"


# 将同步函数包装成异步，以便在 Agents 异步环境中使用
async def async_search_jina(query: str) -> str:
    """异步调用 Jina 搜索"""
    return await asyncio.to_thread(search_jina, query)


async def async_crawl_jina(url: str) -> str:
    """异步调用 Jina 抓取"""
    return await asyncio.to_thread(crawl_jina, url)


def save_report_as_docx(title: str, content: str, filename: str = None):
    """将报告保存为DOCX格式"""
    if not DOCX_AVAILABLE:
        print("DOCX功能不可用，无法保存为Word文档")
        return
    
    if filename is None:
        # 清理文件名中的特殊字符
        clean_title = "".join(c for c in title if c.isalnum() or c in (' ','_','-')).rstrip()
        filename = f"{clean_title}.docx"
    
    try:
        # 创建文档
        doc = Document()
        
        # 添加标题
        heading = doc.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 解析并添加内容
        lines = content.split('\n')
        for line in lines:
            if line.startswith('# '):
                # 一级标题
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                # 二级标题
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                # 三级标题
                doc.add_heading(line[4:], level=3)
            elif line.startswith('- ') or line.startswith('* '):
                # 无序列表
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. '):
                # 有序列表
                doc.add_paragraph(line[3:], style='List Number')
            elif line.strip() == '':
                # 空行，添加段落间距
                doc.add_paragraph()
            else:
                # 普通段落
                paragraph = doc.add_paragraph(line)
                paragraph_format = paragraph.paragraph_format
                paragraph_format.space_after = Pt(6)
        
        # 保存文档
        doc.save(filename)
        print(f"报告已保存为: {filename}")
    except Exception as e:
        print(f"保存DOCX文件时出错: {e}")


def save_report_as_md(title: str, content: str, filename: str = None):
    """将报告保存为Markdown格式"""
    if filename is None:
        # 清理文件名中的特殊字符
        clean_title = "".join(c for c in title if c.isalnum() or c in (' ','_','-')).rstrip()
        filename = f"{clean_title}.md"
    
    try:
        with open(filename, 'w', encoding='utf-8') as f:
            f.write(f"# {title}\n\n")
            f.write(content)
        print(f"报告已保存为: {filename}")
    except Exception as e:
        print(f"保存Markdown文件时出错: {e}")


external_client = AsyncOpenAI(
    api_key=os.environ["OPENAI_API_KEY"],
    base_url="https://dashscope.aliyuncs.com/compatible-mode/v1",
)

# --- 3. 代理定义 (Agents) ---
orchestrator_system_prompt = """
你是一名深度研究专家和项目经理。你的任务是协调整个研究项目，包括：
1. **研究规划 (生成大纲):** 根据用户提供的研究主题和初步搜索结果，生成一个详尽、逻辑严密、结构清晰的报告大纲。大纲必须以严格的 JSON 格式输出，用于指导后续的章节内容检索和起草工作。
2. **报告整合 (组装):** 在所有章节内容起草完成后，将它们整合在一起，形成一篇流畅、连贯、格式优美的最终研究报告。报告必须包括摘要、完整的章节内容、结论和引用来源列表。
"""
DeepResearchAgent = Agent(
    "Deep Research Orchestrator",
    instructions=orchestrator_system_prompt,
    model=OpenAIChatCompletionsModel(
        model="qwen-max",
        openai_client=external_client,
    ),
)

# 3.2. 内容起草代理 (Drafting Agent)
drafting_system_prompt = """
你是一名专业的内容撰稿人。你的任务是将提供的原始网页抓取内容和搜索结果，根据指定的章节主题，撰写成一篇结构合理、重点突出、信息准确的报告章节。
你必须严格遵守以下规则：
1. **聚焦主题:** 严格围绕给定的 '章节主题' 进行撰写。
2. **信息来源:** 只能使用提供的 '原始网页内容' 和 '搜索结果摘要' 中的信息。
3. **格式:** 使用 Markdown 格式。
4. **引用:** 对于文中引用的关键事实和数据，必须在段落末尾用脚注或括号标记引用的来源 URL，例如 [来源: URL]。
"""
DraftingAgent = Agent(
    "Content Drafting Specialist",
    instructions=drafting_system_prompt,
    model=OpenAIChatCompletionsModel(
        model="qwen-max",
        openai_client=external_client,
    ),
)

# 改进
reflection_system__prompt = """
你作为一名专业内容审核员，你需要结合用户的提问和生成的回答，判断回答是否满足了提问的要求。如果满足，则直接回满足，不需要有其他输出。如果不满足，则输出还需要分析的子问题，每行一个子问题。
"""
ReflectionAgent = Agent(
    "Reflection Drafting Specialist",
    instructions=reflection_system__prompt,
    model=OpenAIChatCompletionsModel(
        model="qwen-max",
        openai_client=external_client,
    ),
)

"""
原始的提问 + 生成回答
"""

# 新增：质量评估代理
quality_assessment_prompt = """
你是一名专业的研究报告评审专家。你的任务是对生成的章节内容进行质量评估，并提出改进建议。

评估标准：
1. 内容完整性：是否涵盖了章节主题的所有重要方面
2. 信息准确性：引用的信息是否准确可靠
3. 逻辑结构：段落组织是否清晰，论证是否有力
4. 语言表达：文字是否流畅易懂

请根据以上标准评估章节内容，并提供具体的改进建议。
如果内容质量良好，只需回复"合格"。
如果需要改进，请指出具体问题并提供改进建议，格式如下：
问题1: [具体问题描述]
建议1: [具体改进建议]
问题2: [具体问题描述]
建议2: [具体改进建议]
"""
QualityAssessmentAgent = Agent(
    "Quality Assessment Specialist",
    instructions=quality_assessment_prompt,
    model=OpenAIChatCompletionsModel(
        model="qwen-max",
        openai_client=external_client,
    ),
)

# --- 4. 深度研究核心流程 ---

async def deep_research(query: str, max_sections: int = 5) -> str:
    """
    执行深度研究流程：规划 -> 检索 -> 抓取 -> 起草 -> 整合。
    """
    print(f"\n--- Deep Research for: {query} ---\n")

    # 1. 初步检索
    print("Step 1: 进行初步检索...")
    initial_search_results_str = await async_search_jina(query) # 首次提问 -》 检索结果 （top10的网站的标题和摘要）
    print(initial_search_results_str)

    # 2. 生成研究大纲 (使用 JSON 模式确保结构化输出)
    print("\nStep 2: 基于初步结果生成研究大纲...")

    # 大模型基于主题和初步检索结果，进行章节的规划 （章节标题、章节的关键词）
    init_prompt = f"""研究主题: {query}
初步搜索摘要: {initial_search_results_str}
"""

    outline_prompt = init_prompt + """请根据上述信息，生成一个详细的报告大纲。大纲必须包含一个 'title' 和一个 'sections' 数组。
每个章节对象必须包含 'section_title' 和 'search_keywords' (用于精确检索的关键词)。

示例输出 JSON 格式如下，只要json，不要有其他输出
{
    "title": "关于 XX 的深度研究报告",
    "sections": [
        {"section_title": "引言与背景", "search_keywords": "历史, 现状"},
        {"section_title": "核心要素与机制", "search_keywords": "关键概念, 工作原理"},
        {"section_title": "应用与影响", "search_keywords": "行业应用, 社会影响"}
    ]
}
"""
    try:
        # 调用 Orchestrator Agent 生成 JSON 格式的大纲
        outline_response = await Runner.run(
            DeepResearchAgent,
            outline_prompt,
        )
        print(outline_response)
        outline_json = json.loads(outline_response.final_output.strip("```json").strip("```")) # 解析json

    except Exception as e:
        print(f"Error generating outline: {e}. Falling back to a simple structure.")
        # 失败时提供默认大纲
        outline_json = {
            "title": f"关于 {query} 的深度研究报告",
            "sections": [
                {"section_title": "引言与背景", "search_keywords": f"{query}, 历史, 现状"},
                {"section_title": "核心要素与机制", "search_keywords": f"{query}, 工作原理, 关键技术"},
                {"section_title": "应用与影响", "search_keywords": f"{query}, 行业应用, 社会影响"},
                {"section_title": "结论与展望", "search_keywords": f"{query}, 发展趋势, 挑战"}
            ]
        }

    research_title = outline_json.get("title", f"关于 {query} 的深度研究报告")
    sections = outline_json.get("sections", [])
    if len(sections) > max_sections:
        sections = sections[:max_sections]

    print(f"报告标题: {research_title}")
    print(f"规划了 {len(sections)} 个章节。")

    # 3. 并发处理所有章节
    print("\nStep 3: 并发处理所有章节...")
    
    # 创建所有章节的任务
    section_tasks = []
    for i, section in enumerate(sections):
        task = process_section_with_reflection(section, i+1)
        section_tasks.append(task)
    
    # 并发执行所有章节任务
    drafted_sections = await asyncio.gather(*section_tasks)
    
    # 4. 报告整合与最终输出 (调用 Orchestrator Agent)
    print("\nStep 4: 整合最终研究报告...")
    full_report_draft = "\n\n".join(drafted_sections)

    final_prompt = f"""
    请将以下所有章节内容整合为一篇完整的、专业的深度研究报告。

    **报告标题:** {research_title}

    **已起草的章节内容:**
    {full_report_draft}

    **任务要求:**
    1. 在报告开头添加一个**【摘要】**，总结报告的主要发现和结论。
    2. 保持各章节之间的连贯性。
    3. 在报告末尾添加一个**【结论与展望】**部分（如果大纲中没有）。
    4. 添加一个**【引用来源】**列表，列出所有章节中提到的 URL。
    5. 整体报告必须格式优美，使用 Markdown 格式。
    """

    try:
        final_report = await Runner.run(
            DeepResearchAgent,
            final_prompt,
        )
        return research_title, final_report.final_output
    except Exception as e:
        error_report = f"最终报告整合失败: {e}\n\n已完成的章节草稿:\n{full_report_draft}"
        return research_title, error_report


async def process_section_with_reflection(section: Dict[str, Any], section_index: int) -> str:
    """
    处理单个章节，包含反思和改进机制
    """
    section_title = section.get("section_title")
    search_keywords = section.get("search_keywords")
    print(f"\n--- 处理章节 {section_index}: {section_title} ---")

    # 初始版本生成
    section_content = await generate_section_content(section_title, search_keywords)
    
    # 质量评估和反思循环
    max_iterations = 3  # 最多迭代3次
    for iteration in range(max_iterations):
        print(f"  -> 第 {iteration + 1} 轮质量评估...")
        
        # 评估当前版本
        assessment = await assess_section_quality(section_title, section_content)
        
        # 如果合格则结束循环
        if "合格" in assessment or "满足" in assessment:
            print(f"  -> 章节 {section_title} 通过质量评估")
            break
        
        # 否则根据建议进行改进
        print(f"  -> 根据反馈改进章节内容...")
        section_content = await improve_section_content(
            section_title, search_keywords, section_content, assessment
        )
    
    return f"## {section_title}\n\n{section_content}"


async def generate_section_content(section_title: str, search_keywords: str) -> str:
    """
    生成章节内容的初始版本
    """
    section_query = f"{section_title} 搜索关键词: {search_keywords}"
    section_search_results_str = await async_search_jina(section_query)

    try:
        search_results = json.loads(section_search_results_str)
        urls_to_crawl = [res['url'] for res in search_results if res.get('url')][:2] # top2相关的网页的url
    except:
        print("Warning: Failed to parse search results for crawl.")
        urls_to_crawl = []

    crawled_content = []
    for url in urls_to_crawl:
        content = await async_crawl_jina(url) # 读取最相关网页的内容
        crawled_content.append(f"--- URL: {url} ---\n{content[:3000]}...\n")  # 限制抓取内容长度

    raw_materials = "\n\n".join(crawled_content)

    # 3.3. 内容起草 (调用 Drafting Agent)
    draft_prompt = f"""
    **章节主题:** {section_title}

    **搜索结果摘要:**
    {section_search_results_str[:3000]}... (仅用于参考要点)

    **原始网页内容 (请基于此内容撰写):**
    {raw_materials}

    请根据上述信息，撰写 {section_title} 这一章节的详细内容。
    """

    try:
        section_draft = await Runner.run(
            DraftingAgent,
            draft_prompt,
        )
        return section_draft.final_output
    except Exception as e:
        error_msg = f"章节起草失败: {e}"
        return error_msg


async def assess_section_quality(section_title: str, section_content: str) -> str:
    """
    评估章节内容质量
    """
    assessment_prompt = f"""
    请评估以下章节内容的质量：

    **章节标题:** {section_title}
    
    **章节内容:**
    {section_content}

    请根据内容完整性、信息准确性、逻辑结构和语言表达等方面进行评估。
    如果内容质量良好，只需回复"合格"。
    如果需要改进，请指出具体问题并提供改进建议，格式如下：
    问题1: [具体问题描述]
    建议1: [具体改进建议]
    问题2: [具体问题描述]
    建议2: [具体改进建议]
    """

    try:
        assessment = await Runner.run(
            QualityAssessmentAgent,
            assessment_prompt,
        )
        return assessment.final_output
    except Exception as e:
        print(f"质量评估失败: {e}")
        return "合格"  # 出错时默认通过


async def improve_section_content(section_title: str, search_keywords: str, 
                                current_content: str, assessment: str) -> str:
    """
    根据评估建议改进章节内容
    """
    improvement_prompt = f"""
    请根据以下评估建议改进章节内容：

    **章节标题:** {section_title}
    
    **当前内容:**
    {current_content}
    
    **评估建议:**
    {assessment}
    
    **搜索关键词:**
    {search_keywords}

    请结合评估建议，重新撰写这一章节的内容，确保解决指出的问题并提高整体质量。
    """

    try:
        improved_draft = await Runner.run(
            DraftingAgent,
            improvement_prompt,
        )
        return improved_draft.final_output
    except Exception as e:
        print(f"内容改进失败: {e}")
        return current_content  # 出错时返回原内容


async def main():
    research_topic = "Agentic AI在软件开发中的最新应用和挑战"
    research_title, final_report = await deep_research(research_topic)
    print(final_report)
    
    # 保存为Markdown文件
    save_report_as_md(research_title, final_report)
    
    # 如果可用，也保存为DOCX文件
    if DOCX_AVAILABLE:
        save_report_as_docx(research_title, final_report)


# 使用 Runner 启动异步主函数
if __name__ == "__main__":
    try:
        asyncio.run(main())
    except NameError:
        # Fallback to standard asyncio run if Runner is not defined or preferred
        asyncio.run(main())